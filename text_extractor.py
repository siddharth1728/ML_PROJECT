"""
text_extractor.py - PDF/DOCX Text Extraction & Cleaning Module

Handles ingestion of resume files in PDF and DOCX formats,
extracting raw text and cleaning it through a normalization pipeline.

Cleaning Pipeline:
    1. Remove PDF artifacts (cid:XX character codes)
    2. Normalize unicode characters
    3. Remove non-printable/control characters
    4. Standardize bullet points
    5. Normalize whitespace
"""

import re
import os
import io
import unicodedata
from typing import Union, Optional

from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
from docx import Document


def extract_text_from_pdf(source: Union[str, bytes, io.BytesIO]) -> str:
    """Extract raw text from a PDF file using pdfminer.six."""
    laparams = LAParams(
        line_margin=0.5, word_margin=0.1, char_margin=2.0,
        boxes_flow=0.5, detect_vertical=False
    )
    if isinstance(source, str):
        if not os.path.exists(source):
            raise FileNotFoundError(f"PDF file not found: {source}")
        return extract_text(source, laparams=laparams)
    elif isinstance(source, bytes):
        return extract_text(io.BytesIO(source), laparams=laparams)
    elif isinstance(source, io.BytesIO):
        source.seek(0)
        return extract_text(source, laparams=laparams)
    else:
        raise ValueError(f"Unsupported source type: {type(source)}")


def extract_text_from_docx(source: Union[str, bytes, io.BytesIO]) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    if isinstance(source, str):
        if not os.path.exists(source):
            raise FileNotFoundError(f"DOCX file not found: {source}")
        doc = Document(source)
    elif isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    elif isinstance(source, io.BytesIO):
        source.seek(0)
        doc = Document(source)
    else:
        raise ValueError(f"Unsupported source type: {type(source)}")

    text_parts = []
    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)
    # Extract from tables (resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                text_parts.append(" | ".join(row_texts))
    return "\n".join(text_parts)


def clean_text(raw_text: str) -> str:
    """
    Clean and normalize extracted text through a multi-step pipeline.
    Removes artifacts, normalizes unicode, standardizes whitespace.
    """
    if not raw_text:
        return ""
    text = raw_text
    # Step 1: Remove PDF cid artifacts
    text = re.sub(r'\(cid:\d+\)', '', text)
    # Step 2: Normalize unicode
    text = unicodedata.normalize('NFKD', text)
    # Step 3a: Convert em-dash/en-dash to ASCII ' - ' for downstream parsing
    text = re.sub(r'[\u2013\u2014]', ' - ', text)
    # Step 3b: Remove non-printable chars (keep newlines and tabs)
    text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)
    # Step 4: Standardize bullet-like line starters
    text = re.sub(r'^[\s]*[-*>~]+\s+', '  * ', text, flags=re.MULTILINE)
    # Step 5: Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Step 6: Strip each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    return text.strip()


def extract_text_from_file(
    file_path: Optional[str] = None,
    file_bytes: Optional[bytes] = None,
    filename: Optional[str] = None
) -> str:
    """
    Main entry point for text extraction. Detects format and dispatches
    to the appropriate extractor, then cleans the output.
    """
    if file_path:
        ext = os.path.splitext(file_path)[1].lower()
        source = file_path
    elif file_bytes and filename:
        ext = os.path.splitext(filename)[1].lower()
        source = file_bytes
    else:
        raise ValueError(
            "Either 'file_path' or both 'file_bytes' and 'filename' must be provided."
        )

    if ext == '.pdf':
        raw_text = extract_text_from_pdf(source)
    elif ext == '.docx':
        raw_text = extract_text_from_docx(source)
    else:
        raise ValueError(f"Unsupported file format: '{ext}'. Supported: .pdf, .docx")

    return clean_text(raw_text)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python text_extractor.py <file_path>")
        sys.exit(1)
    fp = sys.argv[1]
    try:
        text = extract_text_from_file(file_path=fp)
        print(f"Extracted {len(text)} chars, {len(text.splitlines())} lines from: {fp}")
        print("=" * 60)
        print(text[:2000])
        if len(text) > 2000:
            print(f"\n... [{len(text) - 2000} more chars truncated]")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)